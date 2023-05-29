from docx import Document

def generate_PSR(project_data, college_name, partner_name):
    # Create a new Word document and add content using the retrieved data
    document = Document()
    documentName = None

    p = project_data # for easier syntaxing

    Name = p['project_name']
    documentName = Name + '_Report.docx'

    document.add_heading('Project Summary Report', level=1)

    document.add_paragraph()
    document.add_paragraph(f"Project: {p['project_name']}")
    document.add_paragraph()
    document.add_paragraph(f"School/College: {college_name}")
    document.add_paragraph()
    document.add_paragraph(f"Partner Institution: {partner_name}")
    document.add_paragraph()
    document.add_paragraph(f"Academic Program: {p['academic_program']}")
    document.add_paragraph()
    document.add_paragraph(f"Objectives of the Project: {p['project_description']}")
    document.add_paragraph()

    # Add paragraphs or tables with the data
    print(documentName)

    # Save the document
    document.save(documentName)
